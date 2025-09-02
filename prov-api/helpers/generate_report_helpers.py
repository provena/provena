import os
import random
from typing import Dict, List, Any, Callable, Set, Optional
from dataclasses import dataclass, field
from enum import Enum

from pydantic import ValidationError 
from config import Config
from fastapi import HTTPException

from ProvenaInterfaces.RegistryAPI import ItemBase, ItemSubType, SeededItem, Node
from helpers.entity_validators import RequestStyle, validate_model_run_id, validate_study_id, UserCipherProxy, ServiceAccountProxy
from helpers.prov_connector import upstream_query, downstream_query
from helpers.registry_helpers import fetch_item_from_registry_with_subtype

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import docx
from docx.shared import RGBColor

BASE_HANDLE_URL = "http://hdl.handle.net/"

class NodeType(str, Enum): 
    INPUTS = "inputs"
    MODEL_RUNS = "model_runs"
    OUTPUTS = "outputs"

class NodeDirection(str, Enum):
    UPSTREAM = "UPSTREAM"
    DOWNSTREAM = "DOWNSTREAM"

@dataclass
class ReportNodeCollection():
    """A collection to manage and organise nodes in a report, grouping them by type: 
    inputs, model runs, and outputs. Each type has a corresponding list to store 
    the nodes and a set to track unique node IDs to silently avoid duplicates.

    Methods: 
        add_node(node: ItemBase, node_type: NodeType) -> None:
            Adds a node to the appropriate category based on its type and subtype.
            Checks whether the node is already present within the collection via
            it's respective set.
    """
    inputs: List[ItemBase] = field(default_factory=list)
    model_runs: List[ItemBase] = field(default_factory=list)
    outputs: List[ItemBase] = field(default_factory=list)
    origin_node: ItemBase | None = None

    # Additional sets to track unique IDs for each category
    input_ids: Set[str] = field(default_factory=set)
    model_run_ids: Set[str] = field(default_factory=set)
    output_ids: Set[str] = field(default_factory=set)

    def add_node(self, node: ItemBase, node_type: NodeType) -> None:

        # Handle nodes of type MODEL_RUN
        if node.item_subtype == ItemSubType.MODEL_RUN and node.id not in self.model_run_ids:
            self.model_runs.append(node)
            self.model_run_ids.add(node.id)
        
        # Handle input nodes (datasets and models)
        elif node_type == NodeType.INPUTS and node.item_subtype in {ItemSubType.MODEL, ItemSubType.DATASET} and node.id not in self.input_ids:
            self.inputs.append(node)
            self.input_ids.add(node.id)
        
        # Handle output nodes (datasets only)
        elif node_type == NodeType.OUTPUTS and node.item_subtype == ItemSubType.DATASET and node.id not in self.output_ids:
            self.outputs.append(node)
            self.output_ids.add(node.id)
        
async def validate_node_id(node_id: str, item_subtype: ItemSubType, request_style: RequestStyle, config: Config) -> ItemBase: 
    """Validates that a provided node (id + subtype) does exist within the registry.
    This only currently works with Model Runs and Study Entities. 

    Parameters
    ----------
    node_id : str
        The ID of the node.
    item_subtype : ItemSubType
        The subtype of the node.
    request_style : RequestStyle
        The style of the request (user, service account etc)
    config : Config
        A config object containing information about the different endpoints of the system.

    Raises
    ------
    HTTPException
        If an invalid subtype is provided.
    HTTPException
        If fetched entity contains an error this exception is raised.
    HTTPException
        If the fetched entity is of type SeededItem this exception is raised.
    """

    try:
        
        validation_mapping: Dict[ItemSubType, Callable] = {
            ItemSubType.MODEL_RUN: validate_model_run_id,
            ItemSubType.STUDY: validate_study_id
        }

        validator_func = validation_mapping.get(item_subtype)

        if not validator_func:
            raise HTTPException(status_code=400, detail=f"Unsupported item subtype for validation.\
                                Received subtype {item_subtype.value}, but the supported subtypes\
                                are {','.join(validation_mapping.keys())}")
        
        response = await validator_func(id=node_id, request_style=request_style, config=config)

        if isinstance(response, str):
            raise HTTPException(status_code=400, detail=f"Validation error with provided {item_subtype.value}: {response}")
        
        if isinstance(response, SeededItem):
            raise HTTPException(status_code=400, detail=f"Seeded item with provided {item_subtype.value} cannot be used for this query!")
        
        return response
    
    except HTTPException as e: 
        raise e
    
    except Exception as e: 
        raise HTTPException(status_code=500, detail=f"Error when validating study or model run entities: {str(e)}")


def parse_nodes(node_list: List[Any]) -> List[Node]: 
    """Parses a list of potential nodes into a list of Valid Node objects.

    This function takes a list of potential nodes, attempts to parse them into 
    the Node Pydantic model, and returns a list of validated Node objects. 

    Parameters
    ----------
    node_list : List[Any]
        A list of items that can potentially be parsed as Node objects.

    Returns
    -------
    List[Node]
        A list of parsed Node objects.

    Raises
    ------
    HTTPException
        Raised if the provided node_list is empty.
    HTTPException
        Raised if validation fails while parsing any node in the list.
    """

    node_list_parsed: List[Node] = []

    # Worst case scenario handling, this is unlikely to happen.
    if len(node_list) == 0:
        raise HTTPException(status_code=400, detail="Cannot process the report generation. Node list is empty.")
    
    try:
        
        for node in node_list:
            try: 
                node_list_parsed.append(Node(**node))
            except ValidationError as e:
                # The node object has not been parsed, so not safe to get it directly. 
                node_id = node.get('id', 'undefined id')
                raise HTTPException(status_code=400, detail=f"Validation error in node with id {node_id}. Error {str(e)}")
            except Exception as e: 
                raise HTTPException(status_code=500, detail=f"Error parsing nodes - {str(e)}")
    
    except HTTPException as e:
        raise e
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected error during node parsing - {str(e)}")
    
    return node_list_parsed


async def generate_report(
        user_cipher: str,
        starting_id: str,
        upstream_depth: int,
        config: Config,
        report_node_collection: ReportNodeCollection
) -> ReportNodeCollection:
    """Generates the report by querying upstream and downstream data, parsing nodes, and filtering results.

    This helper function performs upstream and downstream queries, parses the node responses, filters them 
    for relevant data, and stores the results into a shared dataclass for further report generation.

    Parameters
    ----------
    user : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    starting_id : str
        The starting node ID to perform the queries from.
    upstream_depth : int
        The depth of the upstream query to traverse.
    config : Config
        A config object containing information about the different endpoints of the system.
    report_node_collection : ReportNodeCollection
        A collection to store and manage parsed nodes for report generation.

    Returns
    -------
    ReportNodeCollection
        A collection containing the parsed and filtered nodes, organized by type.
    """

    collection = await fetch_parse_all_upstream_downstream_nodes(
        user_cipher=user_cipher,
        starting_id=starting_id,
        upstream_depth=upstream_depth,
        config=config,
        report_node_collection=report_node_collection
    )

    return collection

async def fetch_parse_all_upstream_downstream_nodes(
    user_cipher: str,
    starting_id: str,
    upstream_depth: int,
    config: Config,
    report_node_collection: ReportNodeCollection,
    downstream_depth: int = 1
) -> ReportNodeCollection:
    
    """
    Fetches and parses nodes in both upstream and downstream directions, and adds them to the collection.

    Parameters
    ----------
    user : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    starting_id : str
        The starting node ID to perform the queries from.
    upstream_depth : int
        The depth of the upstream query to traverse.
    config : Config
        A config object containing information about the different endpoints of the system.
    report_node_collection : ReportNodeCollection
        A collection to store and manage parsed nodes for report generation.
    downstream_depth : int, optional
        The depth of the downstream query to traverse, by default 1.

    Returns
    -------
    ReportNodeCollection
        A collection containing the parsed and filtered nodes from both upstream and downstream queries.

    Raises
    ------
    HTTPException
        If nodes cannot be found in upstream or downstream responses.
    HTTPException
        If any error occurs during fetching and parsing.
    """
    
    try:
        # Fetch both upstream and downstream nodes.
        upstream_response: Dict[str, Any] = upstream_query(starting_id=starting_id, depth=upstream_depth, config=config)
        downstream_response: Dict[str, Any] = downstream_query(starting_id=starting_id, depth=downstream_depth, config=config)

        assert upstream_response.get('nodes'), "Upstream node collections not found!"
        assert downstream_response.get('nodes'), "Downstream node collections not found!"

        nodes_upstream: List[Any] = upstream_response.get('nodes', [])
        nodes_downstream: List[Any] = downstream_response.get('nodes', [])

        # Process the nodes in both directions 
        await process_node_collection(
            nodes = nodes_upstream, 
            direction = NodeDirection.UPSTREAM,
            user_cipher = user_cipher,
            config = config,
            report_node_collection = report_node_collection
        )

        await process_node_collection(
            nodes = nodes_downstream,
            direction = NodeDirection.DOWNSTREAM,
            user_cipher = user_cipher, 
            config = config,
            report_node_collection = report_node_collection
        )

        return report_node_collection
            
    except AssertionError as e: 
        raise HTTPException(status_code=404, detail=f"The provided model run with id {starting_id} - error: {str(e)}")
    
    except Exception as e: 
        raise HTTPException(status_code=500, detail=f"Error fetching upstream/downstream nodes - error: {str(e)}")


async def process_node_collection(
        nodes: List[Any], 
        direction: NodeDirection,
        user_cipher: str, 
        config: Config, 
        report_node_collection: ReportNodeCollection
) -> ReportNodeCollection:
    """
    Processes and validates a collection of nodes, adding them to the appropriate section in the ReportNodeCollection.

    Parameters
    ----------
    nodes : List[Any]
        The list of nodes to process.
    direction : NodeDirection
        The direction of the nodes, either UPSTREAM or DOWNSTREAM.
    user : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    config : Config
        A config object containing information about the different endpoints of the system.
    report_node_collection : ReportNodeCollection
        A collection to store and manage parsed nodes for report generation.

    Returns
    -------
    ReportNodeCollection
        The updated collection with nodes processed and validated for report generation.
    """
    
    parsed_nodes = parse_nodes(nodes)
    node_type = NodeType.INPUTS if direction == NodeDirection.UPSTREAM else NodeType.OUTPUTS

    for node in parsed_nodes:
        if should_load_node(node, direction):
            loaded_item = await fetch_item_from_registry_with_subtype(
                user_cipher=user_cipher, 
                id=node.id,
                item_subtype=node.item_subtype,
                config=config
            )

            #This is typed ignored because we validate the properties of the response in the helper function.
            report_node_collection.add_node(node=loaded_item.item, node_type=node_type) #type:ignore

    return report_node_collection

def should_load_node(node: Node, direction: NodeDirection) -> bool:
    """Determines whether a node should be loaded into the ReportCollection
       and fetches the loaded item if it should.

    Parameters
    ----------
    node : Node
        The node to validate.
    direction : NodeDirection
        The direction of the node (UPSTREAM OR DOWNSTREAM)

    Returns
    -------
    bool
        True if the node should be loaded, False otherwise.
    """
    if direction == NodeDirection.UPSTREAM: 
        return node.item_subtype in{
            ItemSubType.MODEL_RUN,
            ItemSubType.MODEL,
            ItemSubType.DATASET
        }
    
    return node.item_subtype == ItemSubType.DATASET 

def add_hyperlink(paragraph: Paragraph, text: str, url: str, color: Optional[RGBColor] = None) -> Optional[docx.oxml.CT_Hyperlink]:
    # Python-docx does not have a native .add_hyperlink() functionality.
    # Sourced from: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx

    """Adds a hyperlink to a Word document paragraph.

    Parameters
    ----------
    paragraph : Paragraph
        The paragraph to which the hyperlink will be added.
    text : str
        The display text for the hyperlink.
    url : str
        The URL that the hyperlink points to.

    Returns
    -------
    Optional[docx.oxml.CT_Hyperlink]
        The created hyperlink object.
    """

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    #new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.color.rgb = color if color else docx.shared.RGBColor(0, 0, 255) # Default dark-blue color unless specified.
    new_run.font.underline = True  

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_word_file(config: Config, node_collection: ReportNodeCollection) -> str: 
    """
    Generates and saves a temporary Word document with data from a ReportNodeCollection.

    Parameters
    ----------
    node_collection : ReportNodeCollection
        A collection containing nodes to include in the Word document.

    Returns
    -------
    str
        The file path to the newly created Word document.

    Raises
    ------
    HTTPException
        If any error occurs during file generation.
    """

    def format_item_subtype_name(item_subtype: ItemSubType) -> str: 
        """
        Converts an item subtype's name into a more friendly readable format by replacing 
        all underscores and converting all except the first letter into lowercases.

        Parameters
        ----------
        item_subtype : ItemSubType
            The ItemSubType ENUM.

        Returns
        -------
        str
            The formatted or prettified string representing an item subtype.

        Examples
        --------
        format_item_subtype_name(ItemSubType.MODEL_RUN)

        Result: 'Model Run'
        """

        item_subtype_name = item_subtype.name

        # Splits the string by underscore, iterates through each word and capitalises accordingly.
        return ' '.join(word.capitalize() for word in item_subtype_name.split('_'))
    
    try:
        document = Document()
        document.add_heading('Model Run Study Close Out Report', 0)

        # Origin Node Information. 
        origin_node_item_subtype = format_item_subtype_name(node_collection.origin_node.item_subtype)
        display_name = node_collection.origin_node.display_name
        before_colon = f"Report for {origin_node_item_subtype}: "
        after_colon = f"{display_name}"

        # Coloring the origin node information
        heading = document.add_heading(level=1)
        run1 = heading.add_run(before_colon)
        run1.font.color.rgb = RGBColor(0, 0, 255)  # Blue color before colon

        # Add the hyperlink for the after_colon text
        paragraph = heading
        # Adds purple color to the hyperlink
        add_hyperlink(paragraph, text=after_colon, url=BASE_HANDLE_URL + node_collection.origin_node.id, color=RGBColor(128,0,128))

        # Remaining table generation/information.
        table = document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        for row in table.rows:
            row.cells[0].width = Inches(1)  # Narrower left column
            row.cells[1].width = Inches(10)  # Wider right column

        # First row here is the modelling-team-name
        table.cell(0, 0).paragraphs[0].add_run("Modelling Team").bold = True
        cell = table.cell(0,1)
        run = cell.paragraphs[0].add_run("<enter your team name here>")
        run.italic = True

        # Second row here is the inputs
        table.cell(1, 0).paragraphs[0].add_run("Inputs").bold = True
        input_data_cell = table.cell(1,1)
        for input_node in node_collection.inputs:
            paragraph = input_data_cell.add_paragraph()
            # Format the item subtype name in bold.
            bold_run = paragraph.add_run(format_item_subtype_name(input_node.item_subtype))
            bold_run.bold = True

            # Add the rest of the remaining string.
            paragraph.add_run(f": {input_node.display_name}, ")
            add_hyperlink(paragraph, text=input_node.id, url=BASE_HANDLE_URL + input_node.id)
            paragraph.add_run(text = "\n")
        
        # Third row here is the model runs
        table.cell(2, 0).paragraphs[0].add_run("Model Runs").bold = True
        model_run_row_data_cell = table.cell(2,1)
        for model_run_node in node_collection.model_runs:
            paragraph = model_run_row_data_cell.add_paragraph()
            # Format the item subtype name in bold.
            bold_run = paragraph.add_run(format_item_subtype_name(model_run_node.item_subtype))
            bold_run.bold = True

            # Add the rest of the remaining string.
            paragraph.add_run(f": {model_run_node.display_name}, ")
            add_hyperlink(paragraph, text=model_run_node.id, url=BASE_HANDLE_URL + model_run_node.id)
            paragraph.add_run(text = "\n")

        # Fourth row here is the outputs
        table.cell(3, 0).paragraphs[0].add_run("Outputs").bold = True
        output_row_data_cell = table.cell(3,1)
        for output_node in node_collection.outputs:
            paragraph = output_row_data_cell.add_paragraph()
            # Format the item subtype name in bold.
            bold_run = paragraph.add_run(format_item_subtype_name(output_node.item_subtype))
            bold_run.bold = True

            # Add the rest of the remaining string.
            paragraph.add_run(f": {output_node.display_name}, ")
            add_hyperlink(paragraph, text=output_node.id, url=BASE_HANDLE_URL + output_node.id)
            paragraph.add_run(text = "\n")

        file_path = f"{config.TEMP_FILE_LOCATION}/generate_report{str(random.randint(1,100000))}).docx"
        document.save(file_path)

        # Return the file path.
        return file_path
           
    except Exception as e:
        print(f"Error during file generation: {str(e)}") 
        raise HTTPException(status_code=500, detail=f"Error generating the word file: {str(e)}")
    

def remove_file(file_path: str) -> None: 
    """Deletes a specified file within the FAST API server.
    Generic method can be used for any file, in this 
    module used to delete the temporarily created word file.

    Parameters
    ----------
    file_path : str
        The path of the file.

    Raises
    ------
    HTTPException
        File Not Found.
    HTTPException
        Error during deletion of the file.
    """

    try:
        os.remove(file_path)

    except FileNotFoundError as e: 
        raise HTTPException(status_code=500, detail=f"Unable to delete the file with path {file_path} - {str(e)}")
    
    except Exception as e: 
        raise HTTPException(status_code=500, detail=f"Something has gone wrong in deleting the file {str(e)}")

def get_model_runs_from_study(study_node_id:str, config:Config, depth:int = 1) -> List[Node]: 
    """
    Fetches all model runs associated with the respective study.

    Parameters
    ----------
    study_node_id : str
        The ID of the study node.
    config : Config
        A config object containing information about the different endpoints of the system.
    depth : int, optional
        The depth of the downstream query to traverse to, by default 1

    Returns
    -------
    List[Node]
        A list of all model run nodes associated with the study.

    Raises
    ------
    HTTPException
        If no model run nodes are found in the downstream query.
    HTTPException
        If any error occurs during fetching.
    """

    try:
        # Query downstream to get all linked model runs at depth 1. 
        # This may contain more than one model run.
        downstream_model_run_response = downstream_query(starting_id=study_node_id, depth=depth, config=config)
        assert downstream_model_run_response.get('nodes'), f"The study with id {study_node_id} does not have any associated model runs."

        model_run_nodes: List[Node] = parse_nodes(downstream_model_run_response.get('nodes', []))
        return model_run_nodes
    
    except AssertionError as e:
        raise HTTPException(status_code=404, detail=f"{str(e)}")
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching downstream nodes from the provided study with id {study_node_id} - error: {str(e)}")

async def populate_model_run_report(
    user_cipher: str,
    node_id: str,
    upstream_depth: int,
    config: Config,
    report_node_collection: ReportNodeCollection
) -> None:
    """
    A helper function wrapper around the "generate-report" helper and prepares the required parameters.

    Parameters
    ----------
    user : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    node_id : str
        The ID of the model run node.
    upstream_depth : int
        The depth of the upstream query to traverse upto
    config : Config
        A config object containing information about the different endpoints of the system.
    report_node_collection : ReportNodeCollection
        A collection to store and manage parsed nodes for report generation.
    """

    await generate_report(
        user_cipher=user_cipher,
        starting_id=node_id,
        upstream_depth=upstream_depth,
        config=config,
        report_node_collection=report_node_collection
    )

async def populate_study_report(
    user_cipher: str,
    study_node_id:str,
    upstream_depth:int, 
    config:Config, 
    report_node_collection:ReportNodeCollection
) -> None:
    """
    Populates with data from all model runs associated with a study node.
    
    Parameters
    ----------
    user : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    study_node_id : str
        The ID of the study node.
    upstream_depth : int
        The depth of the upstream query to traverse upto
    config : Config
        A config object containing information about the different endpoints of the system.
    report_node_collection : ReportNodeCollection
        A collection to store and manage parsed nodes for report generation.
    """
    
    model_run_nodes = get_model_runs_from_study(study_node_id=study_node_id, config=config)

    # Now branch out with the model runs in here and explore them as above (depth 1-3 upstream and depth 1 downstream)
    # The model run will be the new updated starting id.
    for node in model_run_nodes:
        if node.item_subtype == ItemSubType.MODEL_RUN:
            await populate_model_run_report(
                user_cipher=user_cipher,
                node_id=node.id,
                upstream_depth=upstream_depth,
                config=config,
                report_node_collection=report_node_collection
            )


async def generate_report_helper(
        node_id: str, 
        upstream_depth: int, 
        item_subtype: ItemSubType, 
        config: Config,
        proxy: UserCipherProxy
) -> str:
    
    """
    Generates a study-close-out table for a specified node, supporting nodes with subtypes of model run and study.

    Parameters
    ----------
    node_id : str
        The ID of the node from which the report is generated.
    upstream_depth : int
        The depth of the upstream query to traverse upto from chosen node.
    item_subtype : ItemSubType
        The subtype of the node, indicating whether it's a model run or a study.
    roles : ProtectedRole
        The role of the user making the request, used for making requests to Registry API.
    config : Config
        A config object containing information about the different endpoints of the system.

    Returns
    -------
    str
        The file path to the generated report document.

    Raises
    ------
    HTTPException
        If an unsupported node subtype is requested.
    HTTPException
        Errors in generating word document.
    HTTPException
        HTTP Exceptions caught in other sub-functions.
    HTTPException
        Base Exceptions caught in other sub-functions.
    """

    try:
        # Create the request style, here we assert where the HTTP request came from. 
        request_style = RequestStyle(
            user_direct=None,
            service_account=ServiceAccountProxy(
                user_cipher=proxy.user_cipher,
                direct_service=False
            )
        )

        # Shared dataclass that is passed via reference to other helper functions. 
        # Allows you to store multiple entries due to single instantiation.
        report_nodes: ReportNodeCollection = ReportNodeCollection()

        # do checks accordingly. 
        origin_node: ItemBase = await validate_node_id(
            node_id=node_id,
            item_subtype=item_subtype,
            request_style=request_style,
            config=config
        )

        # Add origin node to dataclass, will be used later in word-doc generation. 
        report_nodes.origin_node = origin_node

        if item_subtype == ItemSubType.MODEL_RUN: 
            await populate_model_run_report(
                user_cipher=proxy.user_cipher, 
                node_id=node_id,
                upstream_depth=upstream_depth,
                config=config,
                report_node_collection=report_nodes
            )
        elif item_subtype == ItemSubType.STUDY: 
            await populate_study_report(
                user_cipher=proxy.user_cipher, 
                study_node_id=node_id,
                upstream_depth=upstream_depth,
                config=config,
                report_node_collection=report_nodes
            )      
        else: 
            raise HTTPException(status_code=400, detail=f"Unsupported node with item subtype {item_subtype.value} requested.\
                                This endpoint only supports subtype of STUDY and MODEL_RUN.")
        
        # All the nodes involved in the model run/study have been populated. 
        generated_doc_path = generate_word_file(config, report_nodes)
        
        if os.path.exists(generated_doc_path): 
            return generated_doc_path
        else: 
            raise HTTPException(status_code=400, detail= "Error generating your study-closeout document")
        
    except HTTPException as e: 
        raise e 
    
    except Exception as e: 
        raise HTTPException(status_code=500, detail=f"Error in report generation: {str(e)}") 